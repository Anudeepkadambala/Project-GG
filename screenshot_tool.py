from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLineEdit, QLabel, QFileDialog, 
                             QRadioButton, QCheckBox, QTextEdit, QMessageBox, QProgressBar)
from PyQt5.QtCore import Qt, pyqtSignal, QThread
from PyQt5.QtGui import QFont, QColor
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import pandas as pd
import os
import time
from PIL import Image
from docx import Document
from docx.shared import RGBColor
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
import imagehash
import re
from urllib.parse import urlparse
import csv
from reportlab.lib.colors import Color
from reportlab.lib.units import inch
from docx.shared import Inches



class WorkerThread(QThread):
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished = pyqtSignal()

    def __init__(self, input_file_path, output_file_path, save_as, capture_login_only):
        super().__init__()
        self.input_file_path = input_file_path
        self.output_file_path = output_file_path
        self.save_as = save_as
        self.capture_login_only = capture_login_only
        self.hash_comparison_logs = []  # Store hash comparison logs for CSV

    def run(self):
        self.process_csv(self.input_file_path, self.output_file_path, self.save_as, self.capture_login_only)
        self.save_hash_comparison_logs_to_csv()

    def capture_screenshot(self, url, screenshot_path):
        CHROME_DRIVER_PATH = 'C:/Driver/chromedriver.exe'  # Set this to your ChromeDriver path
        options = webdriver.ChromeOptions()
        options.add_argument('--disable-gpu')
        options.add_argument('--ignore-certificate-errors')
        options.add_argument('--headless')
        service = Service(CHROME_DRIVER_PATH)
        driver = webdriver.Chrome(service=service, options=options)
        
        try:
            driver.set_window_size(1920, 1080)
            if not url.startswith(('http://', 'https://')):
                raise ValueError("The URL must start with 'http://' or 'https://'")

            driver.get(url)
            time.sleep(15)  # Wait for 15 seconds to ensure the page is loaded

            if self.capture_login_only:
                try:
                    WebDriverWait(driver, 5).until(
                        EC.presence_of_element_located((By.XPATH, "//input[@type='password']"))
                    )
                    driver.save_screenshot(screenshot_path)
                    self.log.emit(f"Saved screenshot for URL: {url}")
                    return True
                except Exception:
                    return False
            else:
                driver.save_screenshot(screenshot_path)
                self.log.emit(f"Saved screenshot for URL: {url}")
                return True

        finally:
            driver.quit()

    def sanitize_filename(self, url):
        return re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', url)

    def color_for_url(self, url):
        """Determine color based on URL scheme and port."""
        parsed_url = urlparse(url)
        if parsed_url.scheme == 'https':
            color = Color(0, 0, 0)  # Black
        elif parsed_url.scheme == 'http':
            color = Color(1, 0, 0)  # Red
        else:
            color = Color(0, 0, 0)  # Black
        
        if parsed_url.port and parsed_url.port not in [80, 443]:
            color = Color(0.82, 0.41, 0.12)  # Chocolate Brown

        return color

    def process_csv(self, input_file_path, output_file_path, save_as, capture_login_only):
        current_dir = os.getcwd()
        df = pd.read_csv(input_file_path)
        urls = df.iloc[:, 3]  # Column D (index 3)

        unique_urls = urls.dropna().drop_duplicates().reset_index(drop=True)
        url_to_hash = {}
        hash_to_urls = {}
        url_to_page = {}

        total_urls = len(unique_urls)
        self.progress.emit(0)

        if save_as.lower() == 'pdf':
            c = canvas.Canvas(output_file_path, pagesize=(1920, 1080))  # Set page size to 1920x1080

            # Add the image as the first page
            c.drawImage('https://news.sophos.com/wp-content/uploads/2024/04/MR-logo.png', 0, 0, width=1920, height=1080)
            c.showPage()

            index_page_content = []

            # Index page
            c.setFont("Helvetica-Bold", 46)
            c.drawCentredString(1920 / 2, 1000, "List of Internet Facing Portals")  # Centered title
            c.setFont("Helvetica", 26)
            table_data = [["URL"]]
            y_position = 970

            current_page = 2  # Initialize current_page here
            for idx, url in enumerate(unique_urls):
                sanitized_url = self.sanitize_filename(url)
                screenshot_path = os.path.join(current_dir, f'screenshot_{sanitized_url}.png')
                
                try:
                    success = self.capture_screenshot(url, screenshot_path)
                    if not success:
                        continue
                    
                    with Image.open(screenshot_path) as img:
                        hash_value = imagehash.average_hash(img)

                    if hash_value not in hash_to_urls:
                        hash_to_urls[hash_value] = [url]
                    else:
                        hash_to_urls[hash_value].append(url)

                    if url in url_to_hash:
                        previous_hash = url_to_hash[url]
                        if previous_hash != hash_value:
                            self.hash_comparison_logs.append([url, previous_hash, hash_value, 'Hash changed'])
                    else:
                        self.hash_comparison_logs.append([url, '', hash_value, 'First entry'])

                    url_to_hash[url] = hash_value

                except Exception:
                    self.hash_comparison_logs.append([url, '', '', 'Error during processing'])
                    continue

                color = self.color_for_url(url)
                index_page_content.append((url, color))
                self.progress.emit(int((idx + 1) / total_urls * 100))

            # Create legend
            legend_data = [
                ['Color Legend', ''],
                ['Black', 'HTTPS URLs'],
                ['Red', 'HTTP URLs'],
                ['Brown', 'URLs with non-standard ports']
            ]

            legend_table = Table(legend_data, colWidths=[3.5*inch, 4.5*inch], rowHeights=[0.5*inch] * len(legend_data))
            legend_table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 18),
                ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
                ('BACKGROUND', (0, 1), (0, 1), colors.black),
                ('TEXTCOLOR', (0, 1), (0, 1), colors.white),
                ('BACKGROUND', (0, 2), (0, 2), colors.red),
                ('TEXTCOLOR', (0, 2), (0, 2), colors.white),
                ('BACKGROUND', (0, 3), (0, 3), colors.Color(0.82, 0.41, 0.12)),  # Chocolate Brown
                ('TEXTCOLOR', (0, 3), (0, 3), colors.white),
            ])

            legend_table.setStyle(legend_table_style)

            # Position the legend table
            table_width, table_height = legend_table.wrap(0, 0)
            x_position = (1920 - table_width) / 2
            y_position = 750
            legend_table.wrapOn(c, 0, 0)
            legend_table.drawOn(c, x_position, y_position)
            

            # Content pages
            c.showPage()
            for hash_value, urls_list in hash_to_urls.items():
                first_url = urls_list[0]
                sanitized_url = self.sanitize_filename(first_url)
                screenshot_path = os.path.join(current_dir, f'screenshot_{sanitized_url}.png')

                with Image.open(screenshot_path) as img:
                    img_width, img_height = img.size

                additional_height = len(urls_list) * 15
                c.setPageSize((img_width, img_height + 50 + additional_height))
                c.drawImage(screenshot_path, 0, 50 + additional_height, width=img_width, height=img_height)

                c.setFont("Helvetica-Bold", 16)
                y_position = additional_height + 30

                color = self.color_for_url(first_url)
                c.setFillColor(color)
                c.drawString(30, y_position, f"URL: {first_url}")
                url_to_page[first_url] = current_page
                y_position -= 15

                for url in urls_list[1:]:
                    color = self.color_for_url(url)
                    c.setFillColor(color)
                    c.drawString(30, y_position, f"Additional URL: {url}")
                    url_to_page[url] = current_page
                    y_position -= 15

                c.showPage()
                current_page += 1

            c.save()
            self.log.emit(f"PDF saved to {output_file_path}")

        if save_as.lower() == 'word':
            doc = Document()
            index_page_content = []

            doc.add_paragraph('List of Internet Facing Portals', style='Title')

            for idx, url in enumerate(unique_urls):
                sanitized_url = self.sanitize_filename(url)
                screenshot_path = os.path.join(current_dir, f'screenshot_{sanitized_url}.png')
                
                try:
                    success = self.capture_screenshot(url, screenshot_path)
                    if not success:
                        continue
                    
                    with Image.open(screenshot_path) as img:
                        hash_value = imagehash.average_hash(img)

                    if hash_value not in hash_to_urls:
                        hash_to_urls[hash_value] = [url]
                    else:
                        hash_to_urls[hash_value].append(url)

                    if url in url_to_hash:
                        previous_hash = url_to_hash[url]
                        if previous_hash != hash_value:
                            self.hash_comparison_logs.append([url, previous_hash, hash_value, 'Hash changed'])
                    else:
                        self.hash_comparison_logs.append([url, '', hash_value, 'First entry'])

                    url_to_hash[url] = hash_value

                except Exception:
                    self.hash_comparison_logs.append([url, '', '', 'Error during processing'])
                    continue

                color = self.color_for_url(url)
                # Convert color to RGB
                rgb_color = RGBColor(int(color.red * 255), int(color.green * 255), int(color.blue * 255))
                index_page_content.append((url, rgb_color))
                paragraph = doc.add_paragraph(url, style='Normal')
                paragraph.runs[0].font.color.rgb = rgb_color

                self.progress.emit(int((idx + 1) / total_urls * 100))

            for hash_value, urls_list in hash_to_urls.items():
                first_url = urls_list[0]
                sanitized_url = self.sanitize_filename(first_url)
                screenshot_path = os.path.join(current_dir, f'screenshot_{sanitized_url}.png')

                if os.path.exists(screenshot_path):
                    # Specify width and height for the image
                    doc.add_picture(screenshot_path, width=Inches(4), height=Inches(3))  # Adjust width and height as needed
                    doc.add_paragraph(f'URL: {first_url}')

                    for url in urls_list[1:]:
                        doc.add_paragraph(f'Additional URL: {url}')

            doc.save(output_file_path)
            self.log.emit(f"Word document saved to {output_file_path}")

    def save_hash_comparison_logs_to_csv(self):
        if self.hash_comparison_logs:
            df = pd.DataFrame(self.hash_comparison_logs, columns=['URL', 'Previous Hash', 'Current Hash', 'Status'])
            df.to_csv('hash_comparison_logs.csv', index=False)


class ScreenshotTool(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        self.setStyleSheet("background-color: #000000;")
        self.setWindowTitle('Screenshot Tool')
        self.setGeometry(100, 100, 800, 500)

        font = QFont("Arial", 12, QFont.Bold)
        label_color = QColor("#e0e0e0")

        input_file_label = QLabel('Input CSV File:')
        input_file_label.setFont(font)
        input_file_label.setStyleSheet(f"color: {label_color.name()};")

        output_file_label = QLabel('Output File Path:')
        output_file_label.setFont(font)
        output_file_label.setStyleSheet(f"color: {label_color.name()};")

        save_as_label = QLabel('Save As:')
        save_as_label.setFont(font)
        save_as_label.setStyleSheet(f"color: {label_color.name()};")

        self.input_file_entry = QLineEdit()
        self.input_file_entry.setStyleSheet("background-color: #444; color: #fff; padding: 5px; border-radius: 5px;")

        self.output_file_entry = QLineEdit()
        self.output_file_entry.setStyleSheet("background-color: #444; color: #fff; padding: 5px; border-radius: 5px;")

        browse_input_button = QPushButton('Browse')
        browse_input_button.setStyleSheet("background-color: #1e88e5; color: #fff; padding: 10px; border-radius: 5px;")
        browse_input_button.clicked.connect(self.browse_input_file)

        browse_output_button = QPushButton('Browse')
        browse_output_button.setStyleSheet("background-color: #1e88e5; color: #fff; padding: 10px; border-radius: 5px;")
        browse_output_button.clicked.connect(self.browse_output_file)

        process_button = QPushButton('Process')
        process_button.setStyleSheet("background-color: #ff5722; color: #fff; padding: 10px; border-radius: 5px;")
        process_button.clicked.connect(self.process_files)

        self.progress_bar = QProgressBar()
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid #444;
                border-radius: 10px;
                background-color: #444;
                text-align: center;
                color: #faf5f5;
                padding: 5px; 
                font-weight: bold;                         
                font-size: 14px;                    
            }
            QProgressBar::chunk {
                background-color: #007A8E;
                border-radius: 10px;
            }
        """)

        self.log_widget = QTextEdit()
        self.log_widget.setStyleSheet("background-color: #222; color: #e0e0e0; padding: 5px; border-radius: 5px;")
        self.log_widget.setReadOnly(True)

        self.save_as_pdf_radio = QRadioButton('PDF')
        self.save_as_pdf_radio.setStyleSheet(f"color: {label_color.name()};")
        self.save_as_word_radio = QRadioButton('Word')
        self.save_as_word_radio.setStyleSheet(f"color: {label_color.name()};")

        self.capture_login_check = QCheckBox('Capture Login Screens Only')
        self.capture_login_check.setStyleSheet(f"color: {label_color.name()};")

        radio_layout = QHBoxLayout()
        radio_layout.addWidget(self.save_as_pdf_radio)
        radio_layout.addWidget(self.save_as_word_radio)

        layout.addWidget(input_file_label)
        layout.addWidget(self.input_file_entry)
        layout.addWidget(browse_input_button)
        layout.addWidget(output_file_label)
        layout.addWidget(self.output_file_entry)
        layout.addWidget(browse_output_button)
        layout.addWidget(save_as_label)
        layout.addLayout(radio_layout)
        layout.addWidget(self.capture_login_check)
        layout.addWidget(self.progress_bar)
        layout.addWidget(process_button)
        layout.addWidget(self.log_widget)

        self.setLayout(layout)

    def browse_input_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, 'Select CSV File', '', 'CSV Files (*.csv)')
        if file_path:
            self.input_file_entry.setText(file_path)

    def browse_output_file(self):
        file_path, _ = QFileDialog.getSaveFileName(self, 'Save Output File', '', 'PDF Files (*.pdf);;Word Files (*.docx)')
        if file_path:
            self.output_file_entry.setText(file_path)

    def process_files(self):
        input_file_path = self.input_file_entry.text()
        output_file_path = self.output_file_entry.text()
        save_as = 'PDF' if self.save_as_pdf_radio.isChecked() else 'Word'
        capture_login_only = self.capture_login_check.isChecked()

        if not input_file_path or not output_file_path:
            QMessageBox.warning(self, 'Warning', 'Please specify both input and output file paths.')
            return

        self.thread = WorkerThread(input_file_path, output_file_path, save_as, capture_login_only)
        self.thread.progress.connect(self.progress_bar.setValue)
        self.thread.log.connect(self.log_widget.append)
        self.thread.finished.connect(self.on_finished)
        self.thread.start()

    def on_finished(self):
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("PPT Killed successfully")
        msg_box.setText("Its done Y'all, check the Output")
        msg_box.setStyleSheet("""
            QLabel {
                color: white;  /* Text color */
            }
            QPushButton {
                color: white;  /* Button text color */
            }
        """)
        msg_box.exec_()

if __name__ == '__main__':
    app = QApplication([])
    tool = ScreenshotTool()
    tool.show()
    app.exec_()
